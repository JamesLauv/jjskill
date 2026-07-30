"""
DOCX内容提取脚本
从Word文件中提取段落文本（含层级标记）和表格内容。

用法: python extract_docx.py <file.docx> [--json]
输出: JSON格式的文档内容（默认）或纯文本
"""
import sys
import json


def extract_docx(filepath):
    from docx import Document
    doc = Document(filepath)

    result = {
        "paragraphs": [],
        "tables": []
    }

    # 提取段落
    for para in doc.paragraphs:
        if para.text.strip():
            style = para.style.name if para.style else ""
            level = 0
            prefix = ""

            if "Heading" in style:
                try:
                    level = int(style.replace("Heading ", ""))
                except ValueError:
                    level = 0
                prefix = "H%d " % level
            elif "List" in style:
                prefix = "- "

            result["paragraphs"].append({
                "text": para.text.strip(),
                "style": style,
                "level": level,
                "prefix": prefix
            })

    # 提取表格
    for ti, table in enumerate(doc.tables, 1):
        table_data = []
        for ri, row in enumerate(table.rows):
            cells = []
            for cell in row.cells:
                # 处理单元格内的换行
                text = cell.text.strip().replace("\n", " ")
                cells.append(text)
            table_data.append({
                "row_index": ri,
                "cells": cells
            })
        result["tables"].append({
            "table_index": ti,
            "rows": table_data
        })

    return result


def print_text(data):
    """纯文本输出格式"""
    import io
    # Windows终端编码兼容
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

    for para in data["paragraphs"]:
        print(f"{para['prefix']}{para['text']}")

    for table in data["tables"]:
        print(f"\n--- Table {table['table_index']} ---")
        for row in table["rows"]:
            print(f"  Row {row['row_index']}: {' | '.join(row['cells'])}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python extract_docx.py <file.docx> [--json]")
        sys.exit(1)

    filepath = sys.argv[1]
    use_json = "--json" in sys.argv

    try:
        data = extract_docx(filepath)
        if use_json:
            import io
            sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
            print(json.dumps(data, ensure_ascii=False, indent=2))
        else:
            print_text(data)
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)
